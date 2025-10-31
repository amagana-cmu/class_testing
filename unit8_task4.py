##### SAME FILE AS TASK3 ##########
###### retrieve_from_gsheet.py ########

import pygsheets
from typing import List, Dict, Any, Optional
import json
from docx import Document
import os


def json_to_excel(json_file_path: str, excel_file_path: str):
    """
    Reads a JSON file produced by scrape_names and writes its contents
    to an Excel file.
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "names"
    
    try:
        # Open and load the JSON file
        with open(json_file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        if not data:
            print(f"Warning: {json_file_path} is empty.")
            wb.save(excel_file_path)
            return
        
        headers = ['name']
        
        # Check the keys of the *first* item to see what columns to add
        # We assume all items in the file have the same structure
        first_item_keys = data[next(iter(data))].keys()
        
        if 'gender' in first_item_keys:
            headers.append('gender')
        if 'usage' in first_item_keys:
            headers.append('usage')
        if 'desc' in first_item_keys:
            headers.append('desc')
            
        # Write the dynamically created header row
        ws.append(headers)
            
        # Iterate over the JSON data
        for name, info in data.items():
            row = [name]
            if 'gender' in headers:
                row.append(info.get('gender'))
            if 'usage' in headers:
                # Join the list of usages into a single comma-separated string
                row.append(', '.join(info.get('usage', [])))
            if 'desc' in headers:
                row.append(info.get('desc'))
            
            # Append the row to the worksheet
            ws.append(row)
            
        # Save the workbook to the specified file
        wb.save(excel_file_path)
        print(f"Successfully converted {json_file_path} to {excel_file_path}")
        
    except FileNotFoundError:
        print(f"Error: The file {json_file_path} was not found.")
    except json.JSONDecodeError:
        print(f"Error: Could not decode JSON from {json_file_path}.")
    except Exception as e:
        print(f"An error occurred during JSON to Excel conversion: {e}")


def csv_to_excel(csv_file_path: str, excel_file_path: str):
    """
    Reads a CSV file and writes its contents to an Excel file.
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "names"
    
    try:
        # Open the CSV file
        with open(csv_file_path, 'r', encoding='utf-8') as f:
            # Use csv.reader to read the file
            reader = csv.reader(f)
            # Iterate over each row in the CSV
            for row in reader:
                # Append the row directly to the worksheet
                ws.append(row)
                
        # Save the workbook
        wb.save(excel_file_path)
        print(f"Successfully converted {csv_file_path} to {excel_file_path}")

    except FileNotFoundError:
        print(f"Error: The file {csv_file_path} was not found.")
    except Exception as e:
        print(f"An error occurred during CSV to Excel conversion: {e}")


def xml_to_excel(xml_file_path: str, excel_file_path: str):
    """
    Reads an XML file produced by scrape_names and writes its contents
    to an Excel file.
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "names"
    
    try:
        # Parse the XML file
        tree = ET.parse(xml_file_path)
        root = tree.getroot()
        
        headers = ['name']
        
        # Find the first <name> element to check its structure
        first_name_element = root.find('name')
        
        if first_name_element is not None:
            # Check for child tags in the *first* element
            if first_name_element.find('gender') is not None:
                headers.append('gender')
            if first_name_element.find('usage') is not None:
                headers.append('usage')
            if first_name_element.find('desc') is not None:
                headers.append('desc')
        else:
            print(f"Warning: {xml_file_path} contains no <name> elements.")
            wb.save(excel_file_path)
            return
            
        # Write the dynamically created header row
        ws.append(headers)
        # --- END FIX ---
        
        # Iterate over all <name> elements in the XML
        for name_elem in root.findall('name'):
            # --- FIX: DYNAMIC ROW CREATION ---
            row = [name_elem.get('value')] # Get 'name' from attribute
            
            if 'gender' in headers:
                gender_tag = name_elem.find('gender')
                row.append(gender_tag.text if gender_tag is not None else None)
            
            if 'usage' in headers:
                usage_tags = name_elem.findall('usage')
                row.append(', '.join([tag.text for tag in usage_tags if tag.text]))
            
            if 'desc' in headers:
                desc_tag = name_elem.find('desc')
                row.append(desc_tag.text if desc_tag is not None else None)
            
            # Append the dynamically created row
            ws.append(row)


        # Save the workbook
        wb.save(excel_file_path)
        print(f"Successfully converted {xml_file_path} to {excel_file_path}")
        
    except FileNotFoundError:
        print(f"Error: The file {xml_file_path} was not found.")
    except ET.ParseError:
        print(f"Error: Could not parse XML from {xml_file_path}.")
    except Exception as e:
        print(f"An error occurred during XML to Excel conversion: {e}")

def get_google_sheet() -> Optional[pygsheets.Worksheet]:
    """
    Authenticates with the Google Sheets API using the service account file
    and returns the 'Accounts' worksheet.
    
    Returns:
        A pygsheets.Worksheet object for the 'Accounts' tab, or None if error.
    """
    try:
        # Authorize using the service account key file specified in the task
        client = pygsheets.authorize(service_file='algebraic-ratio-338019-eb975aa8f686.json')
        
        # Open the Google Spreadsheet by its title, 'Accounts'
        spreadsheet = client.open('Accounts')
        
        # Return the specific worksheet (tab) named 'Accounts'
        worksheet = spreadsheet.worksheet_by_title('Accounts')
        
        return worksheet
        
    except pygsheets.exceptions.SpreadsheetNotFound:
        print("Error: Spreadsheet 'Accounts' not found. Make sure it's shared with the service account email.")
        return None
    except pygsheets.exceptions.WorksheetNotFound:
        print("Error: Worksheet 'Accounts' not found in the spreadsheet.")
        return None
    except FileNotFoundError:
        print("Error: Service account file 'algebraic-ratio-338019-eb975aa8f686.json' not found.")
        return None
    except Exception as e:
        print(f"An error occurred during Google Sheet access: {e}")
        return None

def get_accounts_info(fields: Optional[List[str]] = None, start: int = 0, end: int = 9) -> List[Dict[str, Any]]:
    """
    Retrieves account information from the Google Sheet as a list of dictionaries.
    
    Args:
        fields: A list of column names to retrieve. If None, all columns are retrieved.
        start: The 0-indexed starting row of data to retrieve.
        end: The 0-indexed (inclusive) ending row of data to retrieve.
        
    Returns:
        A list of dictionaries, where each dictionary represents a row of data.
    """
    
    # Call the first function to get the authenticated worksheet object
    worksheet = get_google_sheet()
    
    if worksheet is None:
        print("Failed to retrieve worksheet. Aborting.")
        return [] # Return an empty list if sheet access failed

    try:
        # get_all_records() returns a list of dictionaries,
        # using the first row as the keys. This matches the required output.
        all_records = worksheet.get_all_records(head=1)
        
        # Slice the records based on the 'start' and 'end' parameters.
        # We add +1 to 'end' because Python slicing is exclusive,
        # but the task requirement (and default 0-9) is inclusive.
        sliced_records = all_records[start : end + 1]
        
        # If the 'fields' argument is provided, filter each dictionary
        if fields:
            processed_records = []
            for record in sliced_records:
                # Create a new dictionary containing only the requested fields
                new_record = {}
                for field in fields:
                    if field in record:
                        new_record[field] = record[field]
                processed_records.append(new_record)
            
            return processed_records
        else:
            # If 'fields' is None, return the sliced records as-is
            return sliced_records
            
    except Exception as e:
        print(f"An error occurred while getting account info: {e}")
        return []
    

def generate_reminder(
    first_name: str,
    last_name: str,
    title: str,
    street: str,
    city: str,
    state: str,
    zip_code: str,
    amount_owed: float,
    num_reminders: int,
    output_dir: str
):
    """
    Generates a single Word document reminder letter.
    """
    # Create a new Word document
    document = Document()
    
    # Get the 'No Spacing' style to avoid extra spaces between lines 
    style = document.styles['No Spacing']

    # Add each line as a new paragraph using the 'No Spacing' style
    # Address block 
    document.add_paragraph(f"{title} {first_name} {last_name}", style=style)
    document.add_paragraph(f"{street}", style=style)
    document.add_paragraph(f"{city}, {state} {zip_code}", style=style)

    
    document.add_paragraph(f"Dear {title} {last_name},", style=style)

    # Body of the letter
    # Format the amount to 2 decimal places 
    document.add_paragraph(
        f"Payment on your XYZ account in the amount of ${amount_owed:.2f} is due.",
        style=style
    )

    # Conditional reminder line 
    if num_reminders > 0:
        # Handle plural "reminder(s)"
        reminder_text = "reminders" if num_reminders > 1 else "reminder"
        document.add_paragraph(
            f"We have already sent you {num_reminders} {reminder_text}.",
            style=style
        )

    document.add_paragraph(
        "Please, disregard this notice if you have already made the payment.",
        style=style)

    # Closing 
    document.add_paragraph("Best Regards.", style=style)
    document.add_paragraph("John Doe", style=style)
    document.add_paragraph("Account Manager XYZ", style=style)

    # --- File Naming ---
    # Calculate the new reminder number (num_reminders + 1)
    reminder_num_plus_one = num_reminders + 1
    
    # Create the filename in the specified format 
    filename = f"{last_name}_{first_name}_reminder_{reminder_num_plus_one}.docx"
    
    # Use os.path.join to safely create the full output path
    full_path = os.path.join(output_dir, filename)
    
    # Ensure the output directory exists
    os.makedirs(output_dir, exist_ok=True)

    # Save the document 
    document.save(full_path)
    print(f"Generated reminder for {first_name} {last_name} at {full_path}")

def generate_reminders(start: int, end: int, output_dir: str):
    """
    Retrieves a batch of accounts and generates a reminder for each one.
    """
    print(f"Generating reminders for records {start} to {end}...")
    
    # Get the account data from the Google Sheet 
    accounts = get_accounts_info(start=start, end=end)
    
    if not accounts:
        print("No accounts found to generate reminders for.")
        return

    # Loop through each account dictionary
    for account in accounts:
        try:
            # Call generate_reminder for each account 
            # Map the dictionary keys from the sheet to the function parameters
            generate_reminder(
                first_name=account['First Name'],
                last_name=account['Last Name'],
                title=account['Title'],
                street=account['Street'],
                city=account['City'],
                state=account['State'],
                # Convert ZIP to string to handle all formats
                zip_code=str(account['ZIP']), 
                # Convert Amount Owed to float 
                amount_owed=float(account['Amount Owed']), 
                # [Convert Number of Reminders to int 
                num_reminders=int(account['Number of Reminders']), 
                output_dir=output_dir
            )
        except KeyError as e:
            print(f"Error: Missing expected data {e} for an account. Skipping.")
        except ValueError as e:
            print(f"Error: Invalid data type for an account ({e}). Skipping.")
        except Exception as e:
            print(f"An unexpected error occurred for an account: {e}. Skipping.")
    
    print("Batch reminder generation complete.")


if __name__ == "__main__":
    
    # --- google Test 1: 
    # Retrieve first 3 records (index 0, 1, 2) with all fields
    print("--- Testing get_accounts_info(end=2) ---")
    data_example = get_accounts_info(end=2)
    print(json.dumps(data_example, indent=2))
    
    # --- google Test 2: Retrieve specific fields and a different range ---
    print("\n--- Testing get_accounts_info(fields=..., start=3, end=5) ---")
    specific_fields = ['First Name', 'Last Name', 'Amount Owed']
    data_filtered = get_accounts_info(fields=specific_fields, start=3, end=5)
    print(json.dumps(data_filtered, indent=2))

    # --- google Test 3: Test default parameters ---
    print("\n--- Testing get_accounts_info() (default parameters) ---")
    # Should get rows 0-9, all fields
    data_default = get_accounts_info()
    print(f"Retrieved {len(data_default)} records (default is 10).")
    # Print just the first record to save space
    if data_default:
        print("\nFirst record (default):")
        print(json.dumps(data_default[0], indent=2))    

    # Test 1: Generate a single reminder (like the example) 
    print("\n--- Testing Task 4: generate_reminder (Single File) ---")
    generate_reminder(
        first_name='Rocco', 
        last_name='Hupp', 
        title='Mr.', 
        street='9733 Nancy Wall', 
        city='South Christopherville', 
        state='AL', 
        zip_code='23514', 
        amount_owed=376.60, 
        num_reminders=1, 
        output_dir='test_single_reminder' # Saves to a new folder
    )
    
    # Test 2: Generate a batch of reminders
    print("\n--- Testing Task 4: generate_reminders (Batch Files) ---")
    # This will get the first 3 records (0, 1, 2) and create a doc for each
    generate_reminders(start=0, end=2, output_dir='batch_reminders')
